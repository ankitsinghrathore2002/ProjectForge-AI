from docx import Document


def create_docx(content, filename):

    doc = Document()

    doc.add_heading("BlueprintAI", level=1)

    for line in content.split("\n"):

        doc.add_paragraph(line)

    doc.save(filename)

    return filename